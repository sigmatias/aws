from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent.parent
EVIDENCE_DIR = Path((ROOT / "informe" / "EP3_CURRENT_EVIDENCE_DIR.txt").read_text(encoding="utf-8-sig").strip())
OUT = ROOT / "informe" / "EP3_CUY5132_OYANEDEL_MATIAS.docx"


CAPTURES = [
    (1, "Preparacion del entorno", "Validacion inicial del entorno: identidad AWS, versiones locales, monitor secundario seleccionado y MicroSIP portable disponible."),
    (2, "n8n publicado en AWS", "Dashboard real de n8n accesible por dominio publico nip.io con HTTPS y workflows EP3 cargados."),
    (3, "Flujo n8n Slack + IA", "Workflow n8n con Slack Trigger, AI Agent, modelo OpenRouter y nodo de respuesta Slack."),
    (4, "Flujo REST OpenRouter + Slack", "Workflow n8n con Webhook, llamada REST a OpenRouter, preparacion de respuesta y salida Slack por REST."),
    (5, "Prueba real Slack + IA", "Ejecucion real del webhook n8n: OpenRouter responde y Slack confirma publicacion con ok true, canal, timestamp y bot_id."),
    (6, "Workflow de voz STT + IA + TTS", "Workflow n8n de voz con Webhook Audio, REST ElevenLabs STT, OpenRouter IA, REST ElevenLabs TTS y Respond Audio/Text."),
    (7, "Prueba real audio / STT / IA / TTS", "Evidencia CLI: WAV de entrada, intento real ElevenLabs, STT externo real, IA OpenRouter y archivo MP3 TTS generado."),
    (8, "Issabel Docker en EC2", "Validacion por SSH: EC2 Issabel, contenedor Docker activo, version de Asterisk y extensiones SIP 4001/4002 creadas."),
    (9, "Dashboard Issabel y extensiones", "Interfaz web real de Issabel en AWS mostrando PBX Configuration / Extensions para las extensiones SIP 4001 y 4002."),
    (10, "MicroSIP registrados", "Dos clientes MicroSIP portables configurados como extensiones 4001 y 4002, registrados contra Issabel en AWS."),
    (11, "Asterisk peers online", "Salida real de Asterisk sip show peers con las extensiones 4001 y 4002 registradas en estado OK."),
    (12, "Llamada MicroSIP 4001 -> 4002", "Llamada real contestada entre MicroSIP 4001 y MicroSIP 4002, registrada contra Issabel en AWS."),
    (13, "Canales Asterisk de llamada activa", "core show channels verbose muestra dos canales SIP en estado Up, una llamada activa y BridgeID compartido."),
]


ITEMS = [
    ("Item 1", "Instalacion de n8n en la nube", [1, 2, 3]),
    ("Item 2", "Configuracion del flujo y webhook en n8n", [4, 5]),
    ("Item 3", "Integracion STT/TTS y agente IA", [6, 7]),
    ("Item 4", "Dockerizacion Issabel y llamada entre extensiones", [8, 9, 10, 11, 12, 13]),
]


def shade_cell(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(37, 64, 97)
    return p


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "EEF3F8")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(37, 64, 97)
    p.add_run("\n" + body)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(4)


def add_capture(doc, number, title, description):
    path = EVIDENCE_DIR / f"captura_{number:02d}.png"
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(9.2))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figura {number}. {title}. ")
    r.bold = True
    cap.add_run(description)
    cap.paragraph_format.space_after = Pt(12)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(1.4)
section.right_margin = Cm(1.4)
section.top_margin = Cm(1.2)
section.bottom_margin = Cm(1.2)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10)
styles["Title"].font.name = "Calibri"
styles["Title"].font.size = Pt(24)
styles["Title"].font.bold = True

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("EP3 CUY5132\n")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(37, 64, 97)
subtitle = title.add_run("Integracion con tecnologias emergentes")
subtitle.font.size = Pt(16)
subtitle.font.color.rgb = RGBColor(80, 80, 80)

cover = doc.add_table(rows=5, cols=2)
cover.alignment = WD_TABLE_ALIGNMENT.CENTER
cover.autofit = True
rows = [
    ("Estudiante", "Matias Oyanedel"),
    ("Seccion", "CUY1532-001D"),
    ("Asignatura", "Comunicaciones Unificadas"),
    ("Sigla", "CUY5132"),
    ("Fecha de evidencias", "28 de junio de 2026"),
]
for row, (k, v) in zip(cover.rows, rows):
    shade_cell(row.cells[0], "254061")
    set_cell_text(row.cells[0], k, bold=True, color=(255, 255, 255))
    set_cell_text(row.cells[1], v)

doc.add_paragraph()
intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
intro.add_run(
    "Este informe documenta la implementacion y validacion de la EP3: n8n en AWS, "
    "flujo Slack + IA con API REST, procesamiento de audio/STT/TTS y contenedor Issabel "
    "en Docker con dos extensiones SIP funcionales. Todas las capturas fueron tomadas "
    "desde la segunda pantalla completa y contienen el TXT visible obligatorio con "
    "estudiante, seccion y asignatura."
)

add_note(
    doc,
    "Criterio de evidencia",
    "Se incorporan solo evidencias reales generadas durante la ejecucion. No se incluyen tokens, claves API ni contrasenas en el documento. "
    "Cuando un servicio externo bloqueo una accion por permisos o cuota, se documento la respuesta real y se aplico una contingencia funcional sin simular resultados.",
)

doc.add_page_break()
add_heading(doc, "Resumen de cumplimiento")
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, text in enumerate(["Item", "Objetivo", "Evidencias", "Resultado"]):
    shade_cell(hdr[i], "254061")
    set_cell_text(hdr[i], text, bold=True, color=(255, 255, 255))
for item, objective, nums in ITEMS:
    cells = table.add_row().cells
    set_cell_text(cells[0], item, bold=True)
    set_cell_text(cells[1], objective)
    set_cell_text(cells[2], ", ".join(f"Figura {n}" for n in nums))
    set_cell_text(cells[3], "Validado con capturas y salidas reales")

doc.add_paragraph()
add_bullet(doc, "n8n quedo publicado en AWS mediante dominio publico nip.io y HTTPS.")
add_bullet(doc, "El flujo REST de n8n consulta OpenRouter y publica respuesta real en Slack mediante API.")
add_bullet(doc, "El flujo de audio incluye nodos STT/TTS; ElevenLabs fue probado y bloqueo la cuenta Free Tier, por lo que se documento la contingencia real con STT/TTS alternativo.")
add_bullet(doc, "Issabel corre en Docker sobre EC2, con extensiones 4001 y 4002 registradas y llamada real validada por MicroSIP y Asterisk.")

for item, objective, nums in ITEMS:
    doc.add_page_break()
    add_heading(doc, f"{item}: {objective}")
    if item == "Item 2":
        add_note(
            doc,
            "Nota Slack",
            "El token disponible permitio publicacion real en Slack por conversacion directa del bot. "
            "La publicacion en canal/grupo publico quedo limitada por membresia/permisos del bot, por lo que la evidencia muestra el resultado real de la API sin ocultar esa restriccion.",
        )
    if item == "Item 3":
        add_note(
            doc,
            "Nota TTS/STT",
            "ElevenLabs respondio con bloqueo de Free Tier por actividad inusual. La evidencia conserva ese error real y demuestra continuidad funcional con STT externo y TTS Edge, ademas del workflow n8n con nodos STT/TTS.",
        )
    for n in nums:
        title = next(t for num, t, _ in CAPTURES if num == n)
        desc = next(d for num, _, d in CAPTURES if num == n)
        add_capture(doc, n, title, desc)

doc.add_page_break()
add_heading(doc, "Anexos tecnicos")
annex = doc.add_table(rows=1, cols=2)
annex.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = annex.rows[0].cells
shade_cell(hdr[0], "254061")
shade_cell(hdr[1], "254061")
set_cell_text(hdr[0], "Componente", bold=True, color=(255, 255, 255))
set_cell_text(hdr[1], "Detalle documentado", bold=True, color=(255, 255, 255))
tech_rows = [
    ("n8n AWS", "EC2 ep3-n8n-oyanedel, dominio publico https://3.219.120.166.nip.io, workflows Slack/REST/Audio."),
    ("Slack + REST", "Webhook n8n, API REST OpenRouter con modelo openai/gpt-oss-20b:free y publicacion Slack por chat.postMessage."),
    ("Audio", "Archivo WAV de entrada, transcripcion real, respuesta IA y MP3 TTS generado. Error ElevenLabs documentado."),
    ("Issabel AWS", "EC2 ep3-issabel-oyanedel, IP publica 44.213.63.39, acceso web por puerto 4443, SIP UDP 5060."),
    ("Extensiones", "4001 y 4002 creadas en Issabel y registradas en MicroSIP portable."),
    ("Llamada", "Asterisk mostro dos canales SIP en estado Up, 1 active call y BridgeID compartido durante la prueba."),
]
for k, v in tech_rows:
    cells = annex.add_row().cells
    set_cell_text(cells[0], k, bold=True)
    set_cell_text(cells[1], v)

doc.add_paragraph()
add_heading(doc, "Conclusion", level=2)
conclusion = doc.add_paragraph()
conclusion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
conclusion.add_run(
    "La solucion desplegada cumple el objetivo general de integrar automatizacion n8n, APIs externas de IA, procesamiento de voz y una central Issabel dockerizada. "
    "La validacion mas critica corresponde al item 4: las extensiones 4001 y 4002 fueron registradas contra Issabel y realizaron una llamada real, verificada en MicroSIP y en la CLI de Asterisk."
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
